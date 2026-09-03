"""Personal podcast brief agent. One public episode URL -> a personalized Russian brief."""
from __future__ import annotations

import os
import re
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import imageio_ffmpeg
import bleach
import markdown
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Cm, Pt
from flask import Flask, jsonify, render_template, request, send_file
from openai import OpenAI
from yt_dlp import YoutubeDL

app = Flask(__name__)

SUMMARY_PROMPT = """Ты — персональный агент Даши по смысловой обработке подкастов и интервью.
Не пересказывай выпуск. Отфильтруй его через её реальные цели и стиль мышления.

Контекст: Даша развивает MOOGREEN и сейчас особенно нуждается в практических решениях
для вывода MOOGREEN PRO в реальные B2B-продажи. Ей важны системное мышление, причинные
связи, прозрачное разделение факта, вывода и гипотезы. Она не любит банальности,
самопрезентацию и советы без границ применимости. Приоритет: бизнес, затем саморазвитие,
затем здоровье. Цель — рост прибыли, более быстрые качественные решения, предотвращение
дорогих ошибок и меньшая операционная нагрузка на собственника.

Пиши по-русски, компактно и удобно для первого экрана телефона. Используй ровно этот формат:

[Название / гость]

Главное за 30 секунд
— один наиболее ценный вывод именно для Даши.

Стоит забрать
1. … — почему это важно именно ей.
2. … — почему это важно именно ей.
3. … — почему это важно именно ей.

Применение
— до 3 конкретных действий или решений. Если действий нет: «сохраняем как мыслительную модель, действий сейчас не требуется».

Проверка на прочность
— Факт: только проверяемое или прямо сказанное в выпуске, с указанием, если это утверждение гостя.
— Обоснованная интерпретация: …
— Спорное / требует проверки: …

Детали и таймкоды
— только важные фрагменты. Указывай таймкоды лишь если они есть во входном тексте; не выдумывай.

Вердикт: [обязательно / выборочно / можно пропустить].
"""

# Strict v2: source fidelity first, MOOGREEN implications second.
SUMMARY_PROMPT = """Ты — личный аналитик Даши. Сначала точно объясни разговор, только затем анализируй применимость.
MOOGREEN PRO — самостоятельное B2B-направление для wellness-пространств; не подменяй его розничным каналом.
Не приписывай Даше слова ведущей/гостя. Не выдавай свои планы за слова выпуска. Не выдумывай цифры, сроки, условия, таймкоды или цитаты.

Пиши по-русски, в Markdown, короткими блоками для телефона. Ровно такой порядок:

# [Название / гость]
## Коротко: о чём разговор
3–4 простые строки: кто, какую проблему обсуждают, позиция гостя, итог. Затем: «Читать дальше стоит, если…».
## Краткий лог разговора
6–10 хронологических пунктов: вопрос/тема → ответ, пример, возражение или поворот. Таймкод только если он дан во входе; иначе не ставь.
## Главное из беседы
Таблица Markdown: `Тезис гостя | Основание в разговоре | Статус`.
Статус: `сказано гостем`, `кейс/мнение`, `не доказано`.
## Суть в трёх идеях
Только наиболее ценные идеи из источника.
## Карта решений для MOOGREEN PRO
Таблица: `Идея из выпуска | Применимость | Риск/ограничение | Что проверить`.
Не подразумевай, что описанный канал — правильная стратегия для MOOGREEN.
## Для MOOGREEN: стратегия, гипотезы, предложения
— Стратегический вывод: влияет / не влияет / рано решать — и почему.
— Гипотезы для проверки: до 3.
— Предложения: до 3, каждое помечай `предложение агента`.
— Не делать сейчас: до 3.
— Недостающие данные до решения: до 3.
## Источник и проверяемость
Короткие цитаты лишь дословно и только при уверенности. Укажи, если таймкоды недоступны.
Вердикт: обязательно / выборочно / можно пропустить.
"""


def title_from_url(url: str) -> str:
    host = urlparse(url).netloc.removeprefix("www.")
    return f"Выпуск с {host}"


def direct_mave_audio_url(url: str) -> str | None:
    """Extract the active episode's MP3 from a public Mave episode page."""
    if not urlparse(url).netloc.endswith("mave.digital"):
        return None
    response = requests.get(url, timeout=60, headers={"User-Agent": "Mozilla/5.0"})
    response.raise_for_status()
    match = re.search(
        r"storage/podcasts/[^\"\\]+/episodes/[^\"\\]+\.mp3",
        response.text,
    )
    if not match:
        return None
    return f"https://cdn.mave.digital/{match.group(0)}"


def download_direct_audio(url: str, directory: Path) -> Path:
    target = directory / "source.mp3"
    response = requests.get(url, timeout=120, stream=True)
    response.raise_for_status()
    with target.open("wb") as stream:
        for chunk in response.iter_content(1024 * 1024):
            stream.write(chunk)
    return target


def download_audio(url: str, directory: Path) -> Path:
    """Use yt-dlp for YouTube and podcast pages; direct audio URLs get a safe fallback."""
    mave_audio = direct_mave_audio_url(url)
    if mave_audio:
        return download_direct_audio(mave_audio, directory)
    output = str(directory / "source.%(ext)s")
    options = {
        "outtmpl": output,
        "format": "bestaudio/best",
        "noplaylist": True,
        "quiet": True,
        "no_warnings": True,
        "socket_timeout": 60,
    }
    try:
        with YoutubeDL(options) as ydl:
            info = ydl.extract_info(url, download=True)
            requested = info.get("requested_downloads", [])
            if requested and requested[0].get("filepath"):
                return Path(requested[0]["filepath"])
    except Exception as extraction_error:
        # Some podcast hosts expose a direct MP3 link; downloading it is a valid fallback.
        if re.search(r"\.(mp3|m4a|wav|webm)(?:$|[?&])", url, re.I):
            return download_direct_audio(url, directory)
        raise RuntimeError(
            "Не удалось получить аудио по этой ссылке. Пришли прямую ссылку на выпуск "
            "или попробуй зеркало (YouTube/Spotify)."
        ) from extraction_error
    files = list(directory.glob("source.*"))
    if not files:
        raise RuntimeError("Площадка не отдала аудиофайл.")
    return files[0]


def split_for_transcription(audio_path: Path, directory: Path) -> list[Path]:
    """Make small mono MP3 parts so long episodes fit the transcription upload limit."""
    output = directory / "part-%03d.mp3"
    command = [
        imageio_ffmpeg.get_ffmpeg_exe(), "-y", "-i", str(audio_path), "-vn",
        "-ac", "1", "-ar", "16000", "-b:a", "48k", "-f", "segment",
        "-segment_time", "600", "-reset_timestamps", "1", str(output),
    ]
    completed = subprocess.run(
        command, capture_output=True, text=True, timeout=600, check=False
    )
    if completed.returncode:
        raise RuntimeError("Не удалось подготовить аудио к расшифровке.")
    parts = sorted(directory.glob("part-*.mp3"))
    if not parts:
        raise RuntimeError("Не удалось выделить звук из выпуска.")
    return parts


def transcribe(audio_path: Path, directory: Path) -> str:
    client = OpenAI()
    transcript_parts = []
    for number, part in enumerate(split_for_transcription(audio_path, directory), start=1):
        with part.open("rb") as audio:
            result = client.audio.transcriptions.create(
                model="gpt-transcribe", file=audio, language="ru",
                prompt=(
                    "Русский подкаст. Важные имена и термины: Даша, MOOGREEN, "
                    "MOOGREEN PRO, ВкусВилл, Ozon, Wildberries, B2B."
                ),
            )
        transcript_parts.append(f"[Часть {number}]\n{result.text}")
    return "\n\n".join(transcript_parts)


def make_brief(transcript: str, title: str) -> str:
    client = OpenAI()
    response = client.responses.create(
        model=os.environ.get("SUMMARY_MODEL", "gpt-5-mini"),
        instructions=SUMMARY_PROMPT,
        input=f"Название выпуска: {title}\n\nРасшифровка:\n{transcript}",
    )
    return response.output_text


def render_brief_markdown(brief: str) -> str:
    """Render the model's Markdown while allowing only safe, useful formatting."""
    rendered = markdown.markdown(brief, extensions=["tables", "sane_lists"])
    return bleach.clean(
        rendered,
        tags={
            "h1", "h2", "h3", "p", "ul", "ol", "li", "table", "thead",
            "tbody", "tr", "th", "td", "strong", "em", "blockquote", "code",
            "hr", "br", "a",
        },
        attributes={"a": ["href", "title"]},
        protocols=["http", "https"],
        strip=True,
    )


def add_inline_content(paragraph, element: Tag) -> None:
    """Copy the useful inline formatting from a safe brief HTML fragment."""
    for child in element.children:
        if isinstance(child, NavigableString):
            paragraph.add_run(str(child))
        elif child.name == "br":
            paragraph.add_run().add_break()
        elif child.name in {"strong", "b"}:
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in {"em", "i"}:
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Courier New"
        else:
            add_inline_content(paragraph, child)


def brief_docx(brief_html: str) -> BytesIO:
    """Create an editable Word document from the safe rendered brief."""
    safe_html = bleach.clean(
        brief_html,
        tags={"h1", "h2", "h3", "p", "ul", "ol", "li", "table", "thead", "tbody",
              "tr", "th", "td", "strong", "em", "blockquote", "code", "hr", "br", "a"},
        attributes={"a": ["href", "title"]}, protocols=["http", "https"], strip=True,
    )
    soup = BeautifulSoup(safe_html, "html.parser")
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.7)
    section.left_margin = section.right_margin = Cm(1.8)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for element in soup.contents:
        if isinstance(element, NavigableString) or not isinstance(element, Tag):
            continue
        if element.name in {"h1", "h2", "h3"}:
            level = int(element.name[1])
            document.add_heading(element.get_text(" ", strip=True), level=level)
        elif element.name == "p":
            paragraph = document.add_paragraph()
            add_inline_content(paragraph, element)
        elif element.name == "blockquote":
            paragraph = document.add_paragraph(style="Quote")
            add_inline_content(paragraph, element)
        elif element.name in {"ul", "ol"}:
            style = "List Bullet" if element.name == "ul" else "List Number"
            for item in element.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style=style)
                add_inline_content(paragraph, item)
        elif element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue
            columns = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
            table = document.add_table(rows=0, cols=columns)
            table.style = "Table Grid"
            for row in rows:
                cells = row.find_all(["th", "td"], recursive=False)
                word_cells = table.add_row().cells
                for index, cell in enumerate(cells):
                    add_inline_content(word_cells[index].paragraphs[0], cell)
                    if cell.name == "th":
                        for run in word_cells[index].paragraphs[0].runs:
                            run.bold = True
        elif element.name == "hr":
            document.add_paragraph("—" * 28)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


@app.get("/")
def home():
    return render_template("index.html")


@app.post("/api/brief")
def brief():
    url = (request.json or {}).get("url", "").strip()
    if not url.startswith(("https://", "http://")):
        return jsonify(error="Вставь полную публичную ссылку на эпизод."), 400
    if not os.environ.get("OPENAI_API_KEY"):
        return jsonify(error="В сервисе ещё не задан OPENAI_API_KEY."), 500
    try:
        with tempfile.TemporaryDirectory() as temp:
            audio = download_audio(url, Path(temp))
            transcript = transcribe(audio, Path(temp))
            result = make_brief(transcript, title_from_url(url))
        return jsonify(summary_html=render_brief_markdown(result))
    except Exception as error:  # return a readable operational error, never a stack trace
        return jsonify(error=str(error)), 502


@app.post("/api/export/docx")
def export_docx():
    brief_html = (request.json or {}).get("summary_html", "")
    if not isinstance(brief_html, str) or not brief_html.strip():
        return jsonify(error="Сначала сформируй разбор выпуска."), 400
    if len(brief_html) > 1_500_000:
        return jsonify(error="Разбор слишком большой для экспорта."), 400
    return send_file(
        brief_docx(brief_html),
        as_attachment=True,
        download_name="podcast-brief.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 10000)))
